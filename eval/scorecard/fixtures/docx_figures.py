#!/usr/bin/env python3
"""Generate ``figures.docx``: numbered headings with two captioned figures mid-body.

    uv run --with python-docx python eval/scorecard/fixtures/docx_figures.py [out_dir]

The two PNGs are drawn with ImageMagick ``convert`` (a labelled coloured
rectangle each) into a temporary directory; nothing external is embedded.
Deterministic: fixed core-property timestamps, no author, PNGs stripped of
metadata.
"""

from __future__ import annotations

import subprocess
import sys
import tempfile
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

HERE = Path(__file__).resolve().parent
DEFAULT_OUT = HERE.parents[2] / "tests" / "golden" / "corpus"
FIXED_TIME = datetime(2000, 1, 1, tzinfo=UTC)

FIGURES = (
    ("figure-a.png", "#2b6cb0", "Figure A", "Figure 1: A blue rectangle labelled A, drawn for this fixture."),
    ("figure-b.png", "#c05621", "Figure B", "Figure 2: An orange rectangle labelled B, drawn for this fixture."),
)

PARAGRAPHS = {
    "intro": "This document exists to check that a parser keeps two figures where the author put them: each "
             "figure sits between the paragraph that introduces it and its caption, under a numbered heading.",
    "before_a": "The first figure follows this paragraph. It is a plain blue rectangle with a white label, so "
                "any picture classifier should call it a graphic rather than a photograph.",
    "after_a": "After the first figure the text continues with an unrelated remark about the weather, which "
               "was mild for the season and made the field survey easier than the previous year.",
    "before_b": "The second figure follows this paragraph and is orange. It should not be merged with the first "
                "figure, and its caption must stay below it rather than above.",
    "after_b": "With both figures placed, the closing section sums up what a reader should have seen: two "
               "headings at level one, three at level two, two pictures and two captions.",
    "closing": "Nothing in this document is copied from another source; the prose was written for the fixture.",
}


def draw_png(path: Path, colour: str, label: str) -> None:
    subprocess.run(["convert", "-size", "640x320", f"xc:{colour}", "-gravity", "center", "-fill", "white",
                    "-pointsize", "48", "-annotate", "0", label, "-strip", str(path)],
                   check=True, timeout=60)


def build(pngs: dict[str, Path]) -> Document:
    doc = Document()
    doc.add_heading("Field Notes on Two Figures", level=0)
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(PARAGRAPHS["intro"])
    doc.add_heading("1.1 Scope", level=2)
    doc.add_paragraph("The scope is deliberately narrow: two figures, two captions, one small hierarchy of headings.")
    doc.add_heading("2. The Figures", level=1)
    doc.add_heading("2.1 The First Figure", level=2)
    doc.add_paragraph(PARAGRAPHS["before_a"])
    doc.add_picture(str(pngs["figure-a.png"]), width=Inches(4.5))
    doc.add_paragraph(FIGURES[0][3], style="Caption")
    doc.add_paragraph(PARAGRAPHS["after_a"])
    doc.add_heading("2.2 The Second Figure", level=2)
    doc.add_paragraph(PARAGRAPHS["before_b"])
    doc.add_picture(str(pngs["figure-b.png"]), width=Inches(4.5))
    doc.add_paragraph(FIGURES[1][3], style="Caption")
    doc.add_paragraph(PARAGRAPHS["after_b"])
    doc.add_heading("3. Conclusion", level=1)
    doc.add_paragraph(PARAGRAPHS["closing"])

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.created = FIXED_TIME
    props.modified = FIXED_TIME
    props.title = "Field Notes on Two Figures"
    return doc


def main(argv: list[str]) -> int:
    out_dir = Path(argv[0]) if argv else DEFAULT_OUT
    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / "figures.docx"
    with tempfile.TemporaryDirectory() as scratch:
        pngs: dict[str, Path] = {}
        for name, colour, label, _caption in FIGURES:
            pngs[name] = Path(scratch) / name
            draw_png(pngs[name], colour, label)
        build(pngs).save(str(target))
    print(f"{target} ({target.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
