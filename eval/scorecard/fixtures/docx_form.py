#!/usr/bin/env python3
"""Generate ``form.docx`` and ``form.pdf``: a label/value form with empty cells and checkboxes.

    uv run --with python-docx python eval/scorecard/fixtures/docx_form.py [out_dir]

The form is a two-column table (label, value) with several values left
blank, checkbox rows written with the ballot-box characters U+2610 and
U+2612, and a signature block. The PDF is the docx converted by
``soffice --headless`` so the two fixtures share one source of truth.
"""

from __future__ import annotations

import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
DEFAULT_OUT = HERE.parents[2] / "tests" / "golden" / "corpus"
FIXED_TIME = datetime(2000, 1, 1, tzinfo=timezone.utc)
EMPTY, CHECKED = "☐", "☒"

FIELDS = (
    ("Request number", "EQ-2024-0117"),
    ("Requester", "M. Okafor"),
    ("Department", "Field Survey"),
    ("Equipment", "Total station, tripod, two prisms"),
    ("Serial number", ""),
    ("Loan start", "2024-03-04"),
    ("Loan end", ""),
    ("Priority", f"{CHECKED} Standard   {EMPTY} Urgent   {EMPTY} Emergency"),
    ("Insurance confirmed", f"{EMPTY} Yes   {CHECKED} No"),
    ("Approver", ""),
    ("Notes", ""),
)


def build() -> Document:
    doc = Document()
    doc.add_heading("Equipment Loan Request", level=0)
    doc.add_paragraph("Complete every field. Leave a value blank only when it is not yet known; "
                      "the approver fills in the remaining cells.")
    table = doc.add_table(rows=len(FIELDS) + 1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text, header[1].text = "Field", "Value"
    for cell in header:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for r, (label, value) in enumerate(FIELDS, start=1):
        table.rows[r].cells[0].text = label
        table.rows[r].cells[1].text = value
    doc.add_paragraph()
    doc.add_heading("Declaration", level=1)
    doc.add_paragraph(f"{CHECKED} I confirm the equipment will be returned clean and complete.")
    doc.add_paragraph(f"{EMPTY} I request an extension form to be sent with the equipment.")
    signature = doc.add_paragraph("Signature: ______________________    Date: ____________")
    signature.runs[0].font.size = Pt(11)

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.created = FIXED_TIME
    props.modified = FIXED_TIME
    props.title = "Equipment Loan Request"
    return doc


def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
                   check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=300)
    pdf = out_dir / (docx_path.stem + ".pdf")
    if not pdf.is_file():
        raise SystemExit(f"soffice produced no {pdf}")
    return pdf


def main(argv: list[str]) -> int:
    out_dir = Path(argv[0]) if argv else DEFAULT_OUT
    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / "form.docx"
    build().save(str(target))
    pdf = convert_to_pdf(target, out_dir)
    print(f"{target} ({target.stat().st_size} bytes)")
    print(f"{pdf} ({pdf.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
